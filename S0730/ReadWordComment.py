#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 30/07/2018 22:21
# @Author  : DaBai
# @Site    : 
# @File    : ReadWordComment.py
# @Software: PyCharm

import docx
import pdfkit

file = docx.Document(r"C:\Users\DaBai\Desktop\xiaodu.docx");

print("段落数为："+str(len(file.paragraphs)));


for para in file.paragraphs:
    print(para.text);

for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是:"+file.paragraphs[i].text)


